import streamlit as st
import os
import re
import time
from docx import Document
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LCDocument
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers.ensemble import EnsembleRetriever
from langchain_core.messages import SystemMessage, HumanMessage

st.set_page_config(page_title="Страховой Ассистент РК", page_icon="🛡️")
st.title("🛡️ ИИ-Ассистент по страховому законодательству РК")

GLOSSARY = """
Основные термины страхования в РК:
- Страховой случай — событие, с наступлением которого возникает обязанность страховщика осуществить страховую выплату.
- Страховая премия — сумма денег, которую страхователь уплачивает страховщику за страхование.
- Страховая сумма — сумма, в пределах которой страховщик обязуется осуществить страховую выплату.
- Франшиза — часть убытка, не подлежащая возмещению страховщиком.
- Страхователь — лицо, заключившее договор страхования.
- Страховщик — страховая организация, имеющая лицензию.
"""

with st.sidebar:
    google_api_key = st.text_input("Google API Key", type="password",
                                   help="Ключ из Google AI Studio")
    if google_api_key:
        os.environ["GOOGLE_API_KEY"] = google_api_key

    st.divider()
    st.subheader("📂 База знаний")
    uploaded_files = st.file_uploader(
        "Добавить законы (.docx)",
        type="docx",
        accept_multiple_files=True,
        help="Загрузите новые редакции законов"
    )
    if uploaded_files:
        for f in uploaded_files:
            file_path = os.path.join(".", f.name)
            with open(file_path, "wb") as f_out:
                f_out.write(f.getbuffer())
        st.success(f"✅ {len(uploaded_files)} файл(ов) добавлено. Применится после сброса кэша.")
    
    if st.button("🔄 Сбросить базу знаний", use_container_width=True):
        st.cache_resource.clear()
        st.rerun()

    st.divider()
    st.caption("Поиск по всем загруженным документам. Укажите название закона в вопросе для сужения контекста.")

if not google_api_key:
    st.info("🔑 Введите Google API Key в боковой панели, чтобы начать.")
    st.markdown("""
    ### Как получить ключ:
    1. [Google AI Studio](https://aistudio.google.com/apikey)
    2. Создайте API-ключ (бесплатно, лимит 1500 запросов/день)
    """)
    st.stop()

@st.cache_resource
def load_documents():
    """Читает все .docx файлы и возвращает список документов."""
    files = [f for f in os.listdir('.') if f.endswith('.docx')]
    if not files:
        st.error("В папке нет .docx файлов")
        return []
    all_docs = []
    for f_name in files:
        try:
            doc = Document(f_name)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if not text.strip():
                continue
            articles = re.split(r'\n(?=[Сс]татья\s+\d+)', text)
            for art in articles:
                if art.strip():
                    all_docs.append(LCDocument(
                        page_content=art.strip(),
                        metadata={"source": f_name}
                    ))
            with st.sidebar:
                st.success(f"✅ {f_name}: {len(articles)} статей")
        except Exception as e:
            with st.sidebar:
                st.error(f"❌ Ошибка {f_name}: {e}")
    return all_docs

all_docs = load_documents()
if not all_docs:
    st.error("Документы не загружены. Добавьте .docx файлы в папку проекта или загрузите через интерфейс.")
    st.stop()

@st.cache_resource
def create_retriever(_docs):
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    vectorstore = FAISS.from_documents(_docs, embeddings)
    faiss_retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
    
    bm25_retriever = BM25Retriever.from_documents(_docs, k=5)
    
    ensemble_retriever = EnsembleRetriever(
        retrievers=[faiss_retriever, bm25_retriever],
        weights=[0.6, 0.4]
    )
    return ensemble_retriever

retriever = create_retriever(all_docs)

llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash-lite", temperature=0)

SYSTEM_PROMPT = f"""Ты — высококвалифицированный юрист-эксперт по страховому законодательству Республики Казахстан.
Твоя задача — давать точные и юридически корректные ответы СОГЛАСНО ПРЕДОСТАВЛЕННОМУ КОНТЕКСТУ.

{GLOSSARY}

Правила ответа:
1. Отвечай ТОЛЬКО на основании контекста (документов).
2. Обязательно указывай:
   - Название документа (из поля source)
   - Номер статьи
   - Цитату или пересказ соответствующей нормы
3. Если в контексте нет ответа, скажи: «В предоставленных документах информация не найдена. Рекомендую обратиться к полному тексту закона или проверить актуальную редакцию на adilet.zan.kz.»
4. Не придумывай нормы, отсутствующие в контексте.
5. Отвечай на русском языке, официально-деловым стилем.

Контекст:
{{context}}
"""

if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

if prompt := st.chat_input("Задайте вопрос по страховому законодательству РК..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        with st.spinner("🔍 Анализирую законы..."):
            try:
                docs = retriever.get_relevant_documents(prompt)
                
                max_context_chars = 8000
                context = ""
                for d in docs:
                    chunk = f"Источник: {d.metadata.get('source','')}\n{d.page_content}\n\n"
                    if len(context) + len(chunk) > max_context_chars:
                        chunk = chunk[:max_context_chars - len(context)]
                        context += chunk
                        break
                    context += chunk
                
                system_msg = SystemMessage(content=SYSTEM_PROMPT.replace("{context}", context))
                human_msg = HumanMessage(content=prompt)
                
                answer = llm.invoke([system_msg, human_msg]).content
                
                st.markdown(answer)
                st.session_state.messages.append({"role": "assistant", "content": answer})
                
                with st.expander("📚 Использованные источники (топ-3)"):
                    for i, doc in enumerate(docs[:3], 1):
                        src = doc.metadata.get("source", "неизвестно")
                        st.caption(f"{i}. {src}")
                        st.text(doc.page_content[:400] + "...")
            except Exception as e:
                st.error(f"Ошибка при формировании ответа: {str(e)}")
